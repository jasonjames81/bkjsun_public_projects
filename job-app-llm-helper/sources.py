# sources.py
"""Load applicant background text from a local file path or a web URL.

Self-host model: the app runs on the user's own machine, so reading their local
files (a resume in Documents) and fetching public web links (a published Google Doc,
a personal site) is expected and safe.

SECURITY: do NOT expose this app as a shared multi-tenant server. Server-side URL
fetch would become an SSRF vector and local-path loading would read the host's
filesystem. This feature is intended for the single-user self-host model only.

Supported:
- Local .docx (python-docx) / .pdf (pypdf), parsed in-process — no external
  binaries needed; .txt / .md / .html read directly.
- http(s) URLs: fetched and reduced to readable text. A Google Doc must be
  "Published to the web" (or link-viewable) for its text to come through; private
  docs and most of LinkedIn require login and will not extract.
"""

from __future__ import annotations

import html as html_mod
import io
import re
import tempfile
import urllib.request
from pathlib import Path
from urllib.parse import urlparse

_TIMEOUT = 20
_MAX_BYTES = 5_000_000  # 5 MB cap on loaded/fetched content

# Google Docs "published to the web" page scaffolding to drop.
_GDOCS_BOILERPLATE = (
    "Published using Google Docs",
    "Report abuse",
    "Updated automatically every",
)


class SourceError(Exception):
    """Raised for any expected failure (missing file, bad URL, no text)."""


def _strip_html(raw: str) -> str:
    raw = re.sub(r"(?is)<(script|style)[^>]*>.*?</\1>", " ", raw)
    raw = re.sub(r"(?is)<br\s*/?>", "\n", raw)
    raw = re.sub(r"(?is)</(p|div|li|h[1-6]|tr)>", "\n", raw)
    text = re.sub(r"(?s)<[^>]+>", " ", raw)
    text = html_mod.unescape(text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n[ \t]+", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _strip_gdocs_boilerplate(text: str) -> str:
    lines = text.splitlines()
    last = -1
    for i, line in enumerate(lines[:30]):
        if any(m in line for m in _GDOCS_BOILERPLATE):
            last = i
    return "\n".join(lines[last + 1 :]).lstrip() if last >= 0 else text.lstrip()


def _pdftotext(data: bytes) -> str:
    """Extract text from PDF bytes with pypdf (pure Python — no external binary)."""
    try:
        from pypdf import PdfReader
    except ImportError as e:  # pragma: no cover - dependency is in requirements.txt
        raise SourceError(
            "PDF support needs the `pypdf` package (pip install pypdf)."
        ) from e
    try:
        reader = PdfReader(io.BytesIO(data))
        if reader.is_encrypted:
            # An empty-password unlock covers PDFs that are "encrypted" only to
            # set permissions; a real password can't be recovered here.
            try:
                reader.decrypt("")
            except Exception:  # noqa: BLE001 — fall through to the clear message below
                pass
            if reader.is_encrypted:
                raise SourceError(
                    "that PDF is password-protected. Remove the password (or print "
                    "it to a new PDF), then upload again — or paste the text directly."
                )
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
    except SourceError:
        raise
    except Exception as e:  # noqa: BLE001 — pypdf raises many error types on bad input
        raise SourceError(f"could not read PDF: {e}") from e
    if not text.strip():
        raise SourceError(
            "no selectable text found in that PDF — it may be a scan or image. "
            "Paste the text directly, or use a PDF with real text."
        )
    return text


def _docx_to_text(p: Path) -> str:
    """Extract text from a .docx with python-docx (pure Python — no external binary).

    Pulls body paragraphs plus table-cell text, since résumés often lay out contact
    details and skills in tables.
    """
    try:
        from docx import Document
        from docx.opc.exceptions import PackageNotFoundError
    except ImportError as e:  # pragma: no cover - dependency is in requirements.txt
        raise SourceError(
            "DOCX support needs the `python-docx` package (pip install python-docx)."
        ) from e
    try:
        doc = Document(str(p))
    except PackageNotFoundError as e:
        raise SourceError(
            "that file is not a valid .docx (old .doc files aren't supported — "
            "re-save as .docx, or paste the text directly)."
        ) from e
    lines = [para.text for para in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append("\t".join(cells))
    return "\n".join(lines)


_SUPPORTED_SUFFIXES = (".docx", ".pdf", ".txt", ".md", ".html", ".htm")


def load_path(path: str) -> str:
    p = Path(path).expanduser()
    if not p.exists() or not p.is_file():
        raise SourceError(f"file not found: {p}")
    if p.stat().st_size > _MAX_BYTES:
        raise SourceError("file is too large (over 5 MB)")
    return _parse_file(p)


def load_upload(filename: str, data: bytes) -> str:
    """Parse the bytes of a browser-uploaded file, dispatching on its extension.

    Mirrors load_path but for content that never lands at a user-controlled path —
    the file picker hands us bytes plus a name, so we stage them in a temp file with
    the right suffix and reuse the same extractors.
    """
    if len(data) > _MAX_BYTES:
        raise SourceError("file is too large (over 5 MB)")
    suffix = Path(filename or "").suffix.lower()
    if suffix not in _SUPPORTED_SUFFIXES:
        raise SourceError(
            f"unsupported file type '{suffix or '?'}'. Use .docx, .pdf, .txt, .md, or .html"
        )
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        tmp.write(data)
        tmp_path = Path(tmp.name)
    try:
        return _parse_file(tmp_path)
    finally:
        tmp_path.unlink(missing_ok=True)


def _parse_file(p: Path) -> str:
    suffix = p.suffix.lower()

    if suffix == ".docx":
        return _docx_to_text(p)
    if suffix == ".pdf":
        return _pdftotext(p.read_bytes())
    if suffix in (".txt", ".md"):
        return p.read_text(encoding="utf-8", errors="replace")
    if suffix in (".html", ".htm"):
        return _strip_html(p.read_text(encoding="utf-8", errors="replace"))
    raise SourceError(
        f"unsupported file type '{suffix}'. Use .docx, .pdf, .txt, .md, or .html"
    )


def load_url(url: str) -> str:
    host = urlparse(url).netloc.lower()
    if "linkedin.com" in host:
        # LinkedIn returns HTTP 999 to any non-logged-in fetch — there's no public
        # extraction path. Point the user at the reliable workaround instead.
        raise SourceError(
            "LinkedIn blocks automated fetching. Open your profile, click "
            "“More → Save to PDF”, then upload that PDF here instead."
        )
    req = urllib.request.Request(
        url, headers={"User-Agent": "Mozilla/5.0 (job-app-llm-helper)"}
    )
    try:
        with urllib.request.urlopen(req, timeout=_TIMEOUT) as resp:  # noqa: S310 (self-host only)
            ctype = (resp.headers.get("Content-Type") or "").lower()
            raw = resp.read(_MAX_BYTES + 1)
    except Exception as e:
        raise SourceError(f"could not fetch URL: {e}") from e
    if len(raw) > _MAX_BYTES:
        raise SourceError("remote file is too large (over 5 MB)")

    if "pdf" in ctype or url.lower().split("?")[0].endswith(".pdf"):
        return _pdftotext(raw)
    text = raw.decode("utf-8", errors="replace")
    if "html" in ctype or "<html" in text[:2000].lower():
        text = _strip_gdocs_boilerplate(_strip_html(text))
    return text


# Common paths where an org states its mission / posts news. Tried best-effort on
# top of whatever URL the user pastes; missing ones are skipped silently.
_ORG_PATHS = ("", "/about", "/about-us", "/mission", "/news", "/blog", "/press")
_CRAWL_MAX_CHARS = 40_000


def crawl_site(url: str) -> str:
    """Best-effort fetch of an org's homepage plus common about/news/blog paths.

    Returns concatenated readable text (capped), each section tagged with the URL it
    came from. Self-host only — fetches arbitrary URLs on the host. Raises SourceError
    only if nothing at all could be fetched; partial failures are skipped silently.
    """
    parsed = urlparse(url if "://" in url else f"https://{url}")
    if parsed.scheme not in ("http", "https") or not parsed.netloc:
        raise SourceError("provide a website URL (http/https)")
    base = f"{parsed.scheme}://{parsed.netloc}"
    # Honor a specific path the user gave, then try the common ones.
    candidates: list[str] = []
    given = parsed.path.rstrip("/")
    if given and given != "":
        candidates.append(base + given)
    candidates += [base + p for p in _ORG_PATHS]

    sections: list[str] = []
    total = 0
    seen: set[str] = set()
    for cand in candidates:
        if cand in seen or total >= _CRAWL_MAX_CHARS:
            continue
        seen.add(cand)
        try:
            text = load_url(cand).strip()
        except SourceError:
            continue
        if len(text) < 80:  # skip empty/redirect stubs
            continue
        chunk = text[: _CRAWL_MAX_CHARS - total]
        sections.append(f"--- {cand} ---\n{chunk}")
        total += len(chunk)
    if not sections:
        raise SourceError("couldn't fetch any readable pages from that site")
    return "\n\n".join(sections)


# A bare reference with no scheme that still looks like a host ("linkedin.com/in/me",
# "www.example.org/about") — so we can auto-prepend https:// instead of treating it as
# a local file path and failing with "file not found".
_BARE_URL_RE = re.compile(r"^(?:www\.)?[\w-]+(?:\.[\w-]+)+(?:[/?#].*)?$")


def load_source(ref: str) -> dict:
    """Load text from a file path or URL. Returns {"kind", "text"}; raises SourceError."""
    ref = (ref or "").strip()
    if not ref:
        raise SourceError("provide a file path or URL")
    scheme = urlparse(ref).scheme.lower()
    if scheme in ("http", "https"):
        kind, text = "url", load_url(ref)
    elif scheme in ("file",):
        kind, text = "file", load_path(urlparse(ref).path)
    elif not Path(ref).expanduser().exists() and _BARE_URL_RE.match(ref):
        # No scheme, no such file, but shaped like a domain — treat as a web link.
        kind, text = "url", load_url(f"https://{ref}")
    else:
        kind, text = "file", load_path(ref)
    text = (text or "").strip()
    if not text:
        raise SourceError("no readable text found in that source")
    return {"kind": kind, "text": text}
