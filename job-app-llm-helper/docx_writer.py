# docx_writer.py
"""Build a polished .docx cover letter from the structured model output.

Strips the contact block + signature that the model renders as plain text in the
letter and re-emits them with proper docx formatting. The public edition always
receives the contact dict from the request — there is no contact.json on disk.
"""

from __future__ import annotations

import io
import re
from datetime import date
from profile import _CLOSING_RE, _SALUTATION_RE

from docx import Document
from docx.shared import Pt, RGBColor

_COVER_LETTER_SECTION_RE = re.compile(
    r"##\s*1\.\s*TAILORED COVER LETTER\s*\n(.*?)(?=\n##\s*2\.)",
    re.DOTALL | re.IGNORECASE,
)


def extract_cover_letter_section(structured_output: str) -> str:
    """Pull the cover-letter prose out of the response, dropping any appended sections.

    Legacy multi-section output carried a "## 1. TAILORED COVER LETTER" header; the
    current flow emits the letter alone, optionally followed by a "## 4. EMPLOYER
    APPLICATION QUESTIONS" appendix. Either way, return only the letter — never the
    appended employer Q&A (it shouldn't land in the .docx).
    """
    match = _COVER_LETTER_SECTION_RE.search(structured_output)
    if match:
        return match.group(1).strip()
    return re.split(r"\n##\s", structured_output, maxsplit=1)[0].strip()


def _split_letter_parts(letter_text: str) -> tuple[str, list[str]]:
    """Return (salutation_line, body_paragraphs) from a Claude-rendered letter.

    Strips the contact block + date that precede the salutation, and the closing +
    signature that follow the body. We rebuild those from contact.json.
    """
    lines = letter_text.splitlines()

    # Find salutation line
    salutation_idx = next(
        (i for i, ln in enumerate(lines) if _SALUTATION_RE.match(ln)), -1
    )
    if salutation_idx < 0:
        # No salutation found — treat whole text as body, default salutation
        return "Dear Hiring Manager,", _paragraphs_from_lines(lines)

    salutation = lines[salutation_idx].strip().rstrip(",") + ","

    # Find closing line
    closing_idx = next(
        (
            i
            for i in range(len(lines) - 1, salutation_idx, -1)
            if _CLOSING_RE.match(lines[i])
        ),
        len(lines),
    )

    body_lines = lines[salutation_idx + 1 : closing_idx]
    return salutation, _paragraphs_from_lines(body_lines)


def _paragraphs_from_lines(lines: list[str]) -> list[str]:
    """Group lines into paragraphs separated by blank lines."""
    paragraphs: list[str] = []
    buf: list[str] = []
    for line in lines:
        stripped = line.strip()
        if stripped:
            buf.append(stripped)
        elif buf:
            paragraphs.append(" ".join(buf))
            buf = []
    if buf:
        paragraphs.append(" ".join(buf))
    return [p for p in paragraphs if p]


def build_cover_letter_docx(
    letter_text: str,
    *,
    contact: dict | None = None,
    today: date | None = None,
) -> bytes:
    """Build a formatted cover-letter .docx from raw letter prose. Returns bytes."""
    contact = contact or {}
    today = today or date.today()

    document = Document()

    # Default style: Calibri 11pt
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ---- Header: name + contact line ----
    name_para = document.add_paragraph()
    name_run = name_para.add_run(contact.get("name", ""))
    name_run.bold = True
    name_run.font.size = Pt(14)

    contact_bits = [
        contact.get("city_state", ""),
        contact.get("phone", ""),
        contact.get("email", ""),
    ]
    contact_line = " | ".join(b for b in contact_bits if b)
    if contact_line:
        p = document.add_paragraph()
        r = p.add_run(contact_line)
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    linkedin = contact.get("linkedin", "")
    if linkedin:
        p = document.add_paragraph()
        r = p.add_run(linkedin)
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    document.add_paragraph()  # blank line

    # ---- Date ----
    document.add_paragraph(today.strftime("%B %-d, %Y"))
    document.add_paragraph()  # blank line

    # ---- Letter body ----
    salutation, body_paragraphs = _split_letter_parts(letter_text)
    document.add_paragraph(salutation)
    document.add_paragraph()  # blank line

    for para in body_paragraphs:
        document.add_paragraph(para)
        document.add_paragraph()  # spacer

    # ---- Closing + signature ----
    document.add_paragraph("Sincerely,")
    document.add_paragraph()  # space for signature
    document.add_paragraph()
    document.add_paragraph(contact.get("name", ""))

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_coaching_docx(
    content: str,
    *,
    job_title: str = "",
    org_name: str = "",
) -> bytes:
    """Build a formatted .docx from coaching markdown content. Returns bytes."""
    document = Document()

    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    if org_name or job_title:
        title_parts = [p for p in [org_name, job_title] if p]
        title_para = document.add_paragraph()
        title_run = title_para.add_run(" — ".join(title_parts))
        title_run.bold = True
        title_run.font.size = Pt(14)
        document.add_paragraph()

    for line in content.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("## "):
            p = document.add_paragraph()
            r = p.add_run(stripped[3:].strip())
            r.bold = True
            r.font.size = Pt(12)
        elif stripped.startswith("- "):
            document.add_paragraph(stripped[2:].strip(), style="List Bullet")
        else:
            document.add_paragraph(stripped)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


if __name__ == "__main__":
    # Smoke test: produce a sample docx
    import tempfile
    from pathlib import Path

    sample = """## 1. TAILORED COVER LETTER

Ada Lovelace
London, UK
ada@example.com

May 8, 2026

Dear Hiring Manager,

Thank you for the opportunity to apply. This is a sample paragraph one with several
sentences that span a normal paragraph length to test rendering.

Here is paragraph two, demonstrating that paragraph spacing works correctly.

Sincerely,
Ada Lovelace

## 2. RESUME TAILORING SUGGESTIONS
... (this part should not appear in the docx)
"""
    body = extract_cover_letter_section(sample)
    out = build_cover_letter_docx(
        body,
        contact={
            "name": "Ada Lovelace",
            "city_state": "London, UK",
            "email": "ada@example.com",
        },
    )
    out_path = Path(tempfile.gettempdir()) / "cover_letter_smoke.docx"
    out_path.write_bytes(out)
    print(f"Wrote {out_path} ({len(out):,} bytes)")
