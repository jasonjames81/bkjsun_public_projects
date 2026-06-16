"""File-parsing tests for sources.py — pure Python, no external binaries.

Covers the in-process .docx (python-docx) and .pdf (pypdf) extractors that
replaced the old `pandoc` / `pdftotext` subprocess calls, plus the plain-text and
error paths.
"""

import io
import sys
from pathlib import Path

import pytest

APP_DIR = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(APP_DIR))

import sources  # noqa: E402


def _make_pdf(text: str) -> bytes:
    """Build a minimal valid single-page PDF with one text string, by hand.

    Avoids a binary fixture or a heavyweight PDF-writer dependency; the xref
    offsets are computed so pypdf accepts it.
    """
    objs = [
        b"<</Type/Catalog/Pages 2 0 R>>",
        b"<</Type/Pages/Kids[3 0 R]/Count 1>>",
        b"<</Type/Page/Parent 2 0 R/MediaBox[0 0 300 144]/Contents 4 0 R"
        b"/Resources<</Font<</F1 5 0 R>>>>>>",
        None,
        b"<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>",
    ]
    stream = b"BT /F1 24 Tf 20 60 Td (" + text.encode() + b") Tj ET"
    objs[3] = b"<</Length %d>>stream\n%s\nendstream" % (len(stream), stream)

    out = bytearray(b"%PDF-1.4\n")
    offsets = []
    for i, body in enumerate(objs, 1):
        offsets.append(len(out))
        out += b"%d 0 obj" % i + body + b"endobj\n"
    xref_pos = len(out)
    n = len(objs) + 1
    out += b"xref\n0 %d\n" % n
    out += b"0000000000 65535 f \n"
    for off in offsets:
        out += b"%010d 00000 n \n" % off
    out += b"trailer<</Size %d/Root 1 0 R>>\nstartxref\n%d\n%%%%EOF" % (n, xref_pos)
    return bytes(out)


def _make_docx(paragraphs: list[str]) -> bytes:
    from docx import Document

    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_pdf_upload_extracts_text():
    text = sources.load_upload("resume.pdf", _make_pdf("Resume Sample Text"))
    assert "Resume Sample Text" in text


def test_docx_upload_extracts_paragraphs():
    text = sources.load_upload(
        "resume.docx", _make_docx(["Jane Doe", "Senior Engineer"])
    )
    assert "Jane Doe" in text
    assert "Senior Engineer" in text


def test_docx_upload_extracts_table_cells():
    from docx import Document

    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Email"
    table.rows[0].cells[1].text = "jane@example.com"
    buf = io.BytesIO()
    doc.save(buf)
    text = sources.load_upload("resume.docx", buf.getvalue())
    assert "Email" in text
    assert "jane@example.com" in text


def test_scanned_pdf_with_no_text_raises():
    # A PDF whose only "content" draws nothing extractable -> empty text path.
    empty = _make_pdf(" ")
    with pytest.raises(sources.SourceError, match="no selectable text"):
        sources.load_upload("scan.pdf", empty)


def test_corrupt_pdf_raises_sourceerror():
    with pytest.raises(sources.SourceError):
        sources.load_upload("broken.pdf", b"%PDF-1.4 not really a pdf")


def test_doc_legacy_binary_raises_unsupported():
    # .doc is not in the supported suffixes at all.
    with pytest.raises(sources.SourceError, match="unsupported file type"):
        sources.load_upload("old.doc", b"\xd0\xcf\x11\xe0legacy")


def test_invalid_docx_bytes_raise_sourceerror():
    with pytest.raises(sources.SourceError, match="not a valid .docx"):
        sources.load_upload("fake.docx", b"this is not a zip/docx package")


def test_txt_upload_reads_directly():
    assert "hello world" in sources.load_upload("notes.txt", b"hello world")
