"""Tests for docx_writer.py parsing and rendering functions."""

import sys
from pathlib import Path

import pytest

APP_DIR = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(APP_DIR))

import docx_writer  # noqa: E402
from docx_writer import (  # noqa: E402
    _paragraphs_from_lines,
    _split_letter_parts,
    build_cover_letter_docx,
    extract_cover_letter_section,
)


# ---------------------------------------------------------------------------
# extract_cover_letter_section
# ---------------------------------------------------------------------------


class TestExtractCoverLetterSection:
    def test_legacy_multi_section_output(self):
        text = (
            "## 1. TAILORED COVER LETTER\n\n"
            "Dear Hiring Manager,\n\nI am a strong candidate.\n\n"
            "Sincerely,\nAda\n\n"
            "## 2. RESUME TAILORING SUGGESTIONS\n- emphasize algorithms\n"
        )
        result = extract_cover_letter_section(text)
        assert result.startswith("Dear Hiring Manager,")
        assert "Sincerely," in result
        assert "RESUME TAILORING" not in result

    def test_modern_letter_only_output(self):
        text = "Dear Hiring Manager,\n\nI apply for the role.\n\nSincerely,\nAda"
        result = extract_cover_letter_section(text)
        assert result == text

    def test_letter_with_employer_qa_appendix(self):
        text = (
            "Dear Hiring Manager,\n\nBody.\n\nSincerely,\nAda\n\n"
            "## 4. EMPLOYER APPLICATION QUESTIONS\n\n"
            "Q: Why us?\nA: Because machines.\n"
        )
        result = extract_cover_letter_section(text)
        assert result.startswith("Dear Hiring Manager,")
        assert "EMPLOYER APPLICATION QUESTIONS" not in result
        assert "Why us?" not in result

    def test_bare_text_no_headers(self):
        text = "Just some plain text with no markdown headers at all."
        result = extract_cover_letter_section(text)
        assert result == text


# ---------------------------------------------------------------------------
# _split_letter_parts
# ---------------------------------------------------------------------------


class TestSplitLetterParts:
    def test_normal_letter_with_salutation_and_closing(self):
        text = (
            "Dear Hiring Manager,\n\n"
            "I am excited to apply.\n\n"
            "With great enthusiasm,\n\n"
            "Sincerely,\nAda Lovelace"
        )
        salutation, body = _split_letter_parts(text)
        assert salutation == "Dear Hiring Manager,"
        assert len(body) == 2
        assert "excited to apply" in body[0]
        assert "Sincerely" not in " ".join(body)
        assert "Ada Lovelace" not in " ".join(body)

    def test_no_salutation_defaults_to_hiring_manager(self):
        text = "No greeting here.\n\nJust a paragraph."
        salutation, body = _split_letter_parts(text)
        assert salutation == "Dear Hiring Manager,"
        assert len(body) == 2

    def test_no_closing_takes_all_body(self):
        text = "Dear Hiring Manager,\n\nOnly a body, no closing."
        salutation, body = _split_letter_parts(text)
        assert salutation == "Dear Hiring Manager,"
        assert len(body) == 1
        assert "Only a body" in body[0]

    def test_body_line_containing_sincerely_is_not_closing(self):
        text = (
            "Dear Hiring Manager,\n\n"
            "I am sincere in my belief that sincerely\n"
            "is an overused word in letters.\n\n"
            "Sincerely,\nAda"
        )
        salutation, body = _split_letter_parts(text)
        assert salutation == "Dear Hiring Manager,"
        full_body = " ".join(body)
        # "sincerely" appears as part of body text, not stripped out
        assert "overused word" in full_body


# ---------------------------------------------------------------------------
# _paragraphs_from_lines
# ---------------------------------------------------------------------------


class TestParagraphsFromLines:
    def test_blank_line_separated_paragraphs(self):
        lines = [
            "First paragraph line one.",
            "First paragraph line two.",
            "",
            "Second paragraph.",
        ]
        result = _paragraphs_from_lines(lines)
        assert len(result) == 2
        assert result[0] == "First paragraph line one. First paragraph line two."
        assert result[1] == "Second paragraph."

    def test_single_paragraph(self):
        lines = ["Just one paragraph, no blank lines."]
        result = _paragraphs_from_lines(lines)
        assert len(result) == 1
        assert result[0] == "Just one paragraph, no blank lines."

    def test_whitespace_only_lines_treated_as_blanks(self):
        lines = ["Line A.", "   ", "  \t  ", "Line B."]
        result = _paragraphs_from_lines(lines)
        assert len(result) == 2
        assert result[0] == "Line A."
        assert result[1] == "Line B."

    def test_empty_input(self):
        assert _paragraphs_from_lines([]) == []

    def test_multiple_blank_lines_between(self):
        lines = ["Para one.", "", "", "", "Para two."]
        result = _paragraphs_from_lines(lines)
        assert len(result) == 2


# ---------------------------------------------------------------------------
# build_cover_letter_docx (integration)
# ---------------------------------------------------------------------------


class TestBuildCoverLetterDocx:
    LETTER = (
        "Dear Hiring Manager,\n\n"
        "I am thrilled to apply for the position.\n\n"
        "My background aligns well.\n\n"
        "Sincerely,\nAda Lovelace"
    )
    CONTACT = {
        "name": "Ada Lovelace",
        "city_state": "London, UK",
        "phone": "+44 123 456",
        "email": "ada@example.com",
    }

    def test_returns_valid_docx_bytes(self):
        result = build_cover_letter_docx(self.LETTER, contact=self.CONTACT)
        assert isinstance(result, bytes)
        assert len(result) > 0
        # .docx is a zip archive; check PK header
        assert result[:2] == b"PK"

    def test_contact_name_in_output(self):
        result = build_cover_letter_docx(self.LETTER, contact=self.CONTACT)
        # Parse the docx to inspect text content
        from docx import Document
        import io

        doc = Document(io.BytesIO(result))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Ada Lovelace" in full_text

    def test_email_in_contact_line(self):
        result = build_cover_letter_docx(self.LETTER, contact=self.CONTACT)
        from docx import Document
        import io

        doc = Document(io.BytesIO(result))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "ada@example.com" in full_text

    def test_date_appears_in_output(self):
        from datetime import date

        result = build_cover_letter_docx(
            self.LETTER, contact=self.CONTACT, today=date(2026, 6, 15)
        )
        from docx import Document
        import io

        doc = Document(io.BytesIO(result))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "June 15, 2026" in full_text

    def test_salutation_and_body_in_output(self):
        result = build_cover_letter_docx(self.LETTER, contact=self.CONTACT)
        from docx import Document
        import io

        doc = Document(io.BytesIO(result))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Dear Hiring Manager," in full_text
        assert "thrilled to apply" in full_text

    def test_closing_sincerely_not_duplicated(self):
        """build_cover_letter_docx adds its own Sincerely; the letter's should be stripped."""
        result = build_cover_letter_docx(self.LETTER, contact=self.CONTACT)
        from docx import Document
        import io

        doc = Document(io.BytesIO(result))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        # Should appear exactly once (the one added by build_cover_letter_docx)
        assert full_text.count("Sincerely,") == 1
